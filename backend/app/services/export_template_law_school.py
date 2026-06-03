"""法学院标准格式 · 涉外投资合规法律研究意见书（Word）。"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.config import get_settings
from app.models.scenario import InvestigationScenario
from app.services.disclaimer import DISCLAIMER_FULL_TEXT
from app.services.export_context import export_context, format_export_datetime, safe_export_filename


def _set_run_font(run, name: str = "宋体", size_pt: float = 12, bold: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold)
    return p


def _add_heading(doc: Document, text: str, level: int = 1):
    sizes = {0: 18, 1: 15, 2: 13, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level > 0 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size_pt=sizes.get(level, 12), bold=True)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _decision_label(decision: str) -> str:
    return {"approved": "确认", "rejected": "驳回", "pending": "待复核"}.get(decision, decision)


def _doc_number(scenario_id: int) -> str:
    return f"VELA-{datetime.now().year}-{scenario_id:05d}"


def _cover_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(3.2)
        c1.width = Cm(11.5)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        _set_run_font(r0, size_pt=11, bold=True)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value)
        _set_run_font(r1, size_pt=11)
    doc.add_paragraph()


def build_law_school_docx(scenario: InvestigationScenario) -> tuple[bytes, str]:
    settings = get_settings()
    ctx = export_context(scenario)
    payload = ctx["payload"]
    brief = ctx["brief"]
    review = ctx["review"]
    sections_legal = ctx["sections_legal"]
    scenario_obj = ctx["scenario"]
    now = datetime.now()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    # 页边距（法学院常用）
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    org = settings.export_org_name
    dept = settings.export_org_department
    _add_paragraph(doc, org, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=2)
    _add_paragraph(doc, dept, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=12)

    _add_heading(doc, "涉外投资合规法律研究意见书", level=0)
    _add_paragraph(
        doc,
        f"文号：{_doc_number(scenario_obj.id)}    出具日期：{now.strftime('%Y年%m月%d日')}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        space_after=14,
    )

    sub_sectors = payload.get("detected_sub_sectors") or []
    sub_names = "、".join(s.get("name", "") for s in sub_sectors if s.get("name"))

    _cover_table(
        doc,
        [
            ("致", settings.export_recipient_label),
            ("自", f"{settings.export_org_name} · {settings.export_org_department}"),
            ("关于", f"{scenario_obj.project_name} — 巴西投资合规协查"),
            ("项目地点", f"{scenario_obj.city} · {scenario_obj.state} · {scenario_obj.country}"),
            ("行业专包", payload.get("industry_pack_name") or payload.get("detected_industry_name", scenario_obj.industry)),
            ("动作类型", payload.get("detected_action_type_name", scenario_obj.action_type)),
            ("识别子赛道", sub_names or "—"),
            ("协查状态", scenario_obj.status),
        ],
    )

    _add_paragraph(
        doc,
        "【重要说明】本意见书系基于平台协查流程生成的内部法律研究材料，仅供企业法务部及管理层决策参考，"
        "不构成对巴西法域的正式法律意见，亦不能替代当地执业律师出具的法律文书。",
        size=11,
        space_after=10,
    )

    _add_heading(doc, "一、项目事实概要", level=1)
    _add_paragraph(doc, scenario_obj.description or "（无业务描述）")
    if scenario_obj.investment_structure:
        _add_paragraph(doc, f"投资结构：{scenario_obj.investment_structure}")
    if scenario_obj.employee_count:
        _add_paragraph(doc, f"雇员规模：约 {scenario_obj.employee_count} 人")

    _add_heading(doc, "二、分维度法律核查意见", level=1)
    for section in payload.get("sections", []):
        dim = section.get("dimension_name", "")
        _add_heading(doc, f"（{dim}）", level=2)
        if section.get("description"):
            _add_paragraph(doc, section.get("description", ""), size=11)

        for item in section.get("items", []):
            code = item.get("code", "")
            title = item.get("title", "")
            _add_paragraph(doc, f"{code}  {title}", bold=True, size=12)
            _add_paragraph(doc, item.get("description", ""), size=11)
            if item.get("rationale"):
                _add_paragraph(doc, f"纳入理由：{item.get('rationale')}", size=11)

            legal_item = _find_legal_item(sections_legal, code)
            hits = (legal_item or {}).get("legal_hits") or []
            if hits:
                _add_paragraph(doc, "参考法源：", bold=True, size=11, space_after=2)
                for hit in hits[:3]:
                    line = (
                        f"· {hit.get('title_zh') or hit.get('title_pt')} "
                        f"（{hit.get('source_label', '')}，匹配度 {hit.get('match_score', 0)}）"
                    )
                    _add_paragraph(doc, line, size=11, space_after=2)
            else:
                _add_paragraph(doc, "· 未检索到可引用法源片段，建议外聘当地律师补充。", size=11)

            brief_item = _find_brief_item(brief, code)
            if brief_item:
                _add_paragraph(doc, "风险说明（中文）：", bold=True, size=11, space_after=2)
                _add_paragraph(doc, brief_item.get("risk_zh", ""), size=11)
                if brief_item.get("risk_pt"):
                    _add_paragraph(doc, "Nota de risco (Português)：", bold=True, size=11, space_after=2)
                    _add_paragraph(doc, brief_item.get("risk_pt", ""), size=11)

    _add_heading(doc, "三、综合结论与风险提示", level=1)
    if brief:
        _add_paragraph(doc, brief.get("summary_zh", ""))
        if brief.get("summary_pt"):
            _add_paragraph(doc, brief.get("summary_pt", ""), size=11)
        blocked = brief.get("blocked_items") or []
        if blocked:
            _add_paragraph(doc, "需进一步复核事项：", bold=True)
            for item in blocked:
                _add_paragraph(
                    doc,
                    f"· {item.get('code')} — {item.get('title')}（{item.get('block_reason', '')}）",
                    size=11,
                )
    else:
        _add_paragraph(doc, "（简报尚未生成，请返回平台完成法源绑定与简报汇编。）")

    if review and review.get("items"):
        _add_heading(doc, "四、法务复核意见", level=1)
        _add_paragraph(
            doc,
            f"复核人：{review.get('reviewer_name', '—')}    "
            f"复核状态：{review.get('status', '—')}    "
            f"定稿时间：{format_export_datetime(review.get('finalized_at'))}",
            size=11,
        )
        for item in review.get("items", []):
            decision = _decision_label(item.get("decision", "pending"))
            line = f"{item.get('code')} — {item.get('title')}：{decision}"
            if item.get("external_counsel_required"):
                line += " 【需外聘当地律所】"
            if item.get("comment"):
                line += f"（批注：{item.get('comment')}）"
            _add_paragraph(doc, f"· {line}", size=11, space_after=3)

    _add_heading(doc, "附录 · 免责声明", level=1)
    for para in DISCLAIMER_FULL_TEXT.split("\n\n"):
        _add_paragraph(doc, para, size=10.5)

    _add_paragraph(doc, "", space_after=18)
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["经办律师/法务", "复核律师/法务负责人", "出具日期"]
    values = [
        review.get("reviewer_name", "____________") if review else "____________",
        "____________",
        format_export_datetime(review.get("finalized_at") if review else None)
        if review and review.get("finalized_at")
        else now.strftime("%Y年%m月%d日"),
    ]
    for col, (h, v) in enumerate(zip(headers, values)):
        rh = sig_table.rows[0].cells[col].paragraphs[0].add_run(h)
        _set_run_font(rh, size_pt=10.5, bold=True)
        rv = sig_table.rows[1].cells[col].paragraphs[0].add_run(str(v)[:40])
        _set_run_font(rv, size_pt=10.5)

    buf = io.BytesIO()
    doc.save(buf)
    filename = safe_export_filename(scenario_obj.project_name, "法律研究意见书.docx")
    return buf.getvalue(), filename


def _find_legal_item(sections_legal: list[dict[str, Any]], code: str) -> dict[str, Any] | None:
    for section in sections_legal:
        for item in section.get("items", []):
            if item.get("code") == code:
                return item
    return None


def _find_brief_item(brief: dict[str, Any], code: str) -> dict[str, Any] | None:
    for section in brief.get("sections", []):
        for item in section.get("items", []):
            if item.get("code") == code:
                return item
    return None
