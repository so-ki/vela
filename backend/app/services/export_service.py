from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.core.config import get_settings
from app.models.scenario import InvestigationScenario
from app.services.disclaimer import DISCLAIMER_FULL_TEXT
from app.services.export_context import export_context, format_export_datetime, safe_export_filename
from app.services.export_template_law_school import build_law_school_docx

pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def _pdf_style(name: str, **kwargs) -> ParagraphStyle:
    base = {"fontName": "STSong-Light", "fontSize": 11, "leading": 16}
    base.update(kwargs)
    return ParagraphStyle(name, **base)


def _pdf_escape(text: str) -> str:
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_docx(scenario: InvestigationScenario) -> tuple[bytes, str]:
    settings = get_settings()
    if settings.export_template == "law_school":
        return build_law_school_docx(scenario)
    return build_legacy_docx(scenario)


def build_sample_docx(scenario: InvestigationScenario) -> tuple[bytes, str]:
    """Backward-compatible alias."""
    return build_docx(scenario)


def build_legacy_docx(scenario: InvestigationScenario) -> tuple[bytes, str]:
    ctx = export_context(scenario)
    payload = ctx["payload"]
    brief = ctx["brief"]
    review = ctx["review"]
    scenario_obj = ctx["scenario"]
    sections_legal = ctx["sections_legal"]
    checklist = payload

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(11)

    title = doc.add_heading("拉美投资合规协查底稿", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"项目名称：{scenario_obj.project_name}")
    doc.add_paragraph(f"地点：{scenario_obj.city} · {scenario_obj.state} · {scenario_obj.country}")
    doc.add_paragraph(f"协查法域：{checklist.get('industry_pack_name') or checklist.get('detected_industry_name', scenario_obj.industry)}")
    sub_sectors = checklist.get("detected_sub_sectors") or []
    if sub_sectors:
        names = "、".join(s.get("name", "") for s in sub_sectors if s.get("name"))
        if names:
            doc.add_paragraph(f"识别子赛道：{names}")
    doc.add_paragraph(f"动作类型：{checklist.get('detected_action_type_name', scenario_obj.action_type)}")
    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"场景状态：{scenario_obj.status}")

    doc.add_heading("一、业务场景摘要", level=1)
    doc.add_paragraph(scenario_obj.description)
    if scenario_obj.investment_structure:
        doc.add_paragraph(f"投资结构：{scenario_obj.investment_structure}")
    if scenario_obj.investment_destination:
        doc.add_paragraph(f"投资目的地：{scenario_obj.investment_destination}")
    if scenario_obj.funding_source:
        doc.add_paragraph(f"资金来源：{scenario_obj.funding_source}")
    if scenario_obj.project_content_scale:
        doc.add_paragraph(f"主要内容：{scenario_obj.project_content_scale}")
    if scenario_obj.known_risks:
        doc.add_paragraph(f"已知风险与关注事项：{scenario_obj.known_risks}")
    if scenario_obj.employee_count:
        doc.add_paragraph(f"雇员规模：约 {scenario_obj.employee_count} 人")

    doc.add_heading("二、专项核查清单", level=1)
    for section in payload.get("sections", []):
        doc.add_heading(section.get("dimension_name", ""), level=2)
        doc.add_paragraph(section.get("description", ""))
        for item in section.get("items", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{item.get('code')}] {item.get('title')} ").bold = True
            p.add_run(f"（{item.get('priority', '')} 优先级）")
            doc.add_paragraph(item.get("description", ""))

    doc.add_heading("三、法源绑定摘要", level=1)
    for section in sections_legal:
        for item in section.get("items", []):
            hits = item.get("legal_hits") or []
            if not hits:
                continue
            doc.add_paragraph(f"{item.get('code')} — {item.get('title')}", style="List Bullet")
            for hit in hits[:2]:
                doc.add_paragraph(
                    f"  · {hit.get('title_zh') or hit.get('title_pt')} "
                    f"(匹配度 {hit.get('match_score', 0)}，{hit.get('source_label', '')})"
                )

    if brief:
        doc.add_heading("四、中葡双语法律风险简报", level=1)
        doc.add_heading("执行摘要（中文）", level=2)
        doc.add_paragraph(brief.get("summary_zh", ""))
        doc.add_heading("Resumo Executivo (Português)", level=2)
        doc.add_paragraph(brief.get("summary_pt", ""))

        for section in brief.get("sections", []):
            doc.add_heading(section.get("dimension_name", ""), level=2)
            doc.add_paragraph(section.get("summary_zh", ""))
            for item in section.get("items", []):
                if item.get("gate_status") != "passed":
                    continue
                doc.add_paragraph(f"{item.get('code')} {item.get('title')}", style="List Bullet")
                doc.add_paragraph(item.get("risk_zh", ""))
                doc.add_paragraph(item.get("risk_pt", ""))

        blocked = brief.get("blocked_items") or []
        if blocked:
            doc.add_heading("需法务复核条目", level=2)
            for item in blocked:
                doc.add_paragraph(
                    f"{item.get('code')} — {item.get('title')}（{item.get('block_reason', '')}）"
                )

    if review and review.get("items"):
        doc.add_heading("五、法务复核记录", level=1)
        doc.add_paragraph(
            f"复核人：{review.get('reviewer_name', '')} · 状态：{review.get('status', '')}"
        )
        if review.get("finalized_at"):
            doc.add_paragraph(f"定稿时间：{format_export_datetime(review.get('finalized_at'))}")
        for item in review.get("items", []):
            decision = {"approved": "确认", "rejected": "驳回", "pending": "待复核"}.get(
                item.get("decision", "pending"), item.get("decision")
            )
            line = f"{item.get('code')} — {item.get('title')}：{decision}"
            if item.get("comment"):
                line += f"（批注：{item.get('comment')}）"
            doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("六、免责声明", level=1)
    for para in DISCLAIMER_FULL_TEXT.split("\n\n"):
        doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    filename = safe_export_filename(scenario_obj.project_name, "协查底稿.docx")
    return buf.getvalue(), filename


def build_sample_pdf(scenario: InvestigationScenario) -> tuple[bytes, str]:
    ctx = export_context(scenario)
    payload = ctx["payload"]
    brief = ctx["brief"]
    review = ctx["review"]
    scenario_obj = ctx["scenario"]
    sections_legal = ctx["sections_legal"]

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm)
    h1 = _pdf_style("H1", fontSize=16, leading=22, spaceAfter=8)
    h2 = _pdf_style("H2", fontSize=13, leading=18, spaceBefore=8, spaceAfter=4)
    body = _pdf_style("Body")
    story: list[Any] = []

    story.append(Paragraph(_pdf_escape("拉美投资合规协查底稿"), h1))
    story.append(Paragraph(_pdf_escape(f"项目名称：{scenario_obj.project_name}"), body))
    story.append(Paragraph(_pdf_escape(f"地点：{scenario_obj.city} · {scenario_obj.state} · {scenario_obj.country}"), body))
    story.append(Paragraph(_pdf_escape(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}"), body))
    story.append(Spacer(1, 6))

    story.append(Paragraph(_pdf_escape("一、业务场景摘要"), h2))
    story.append(Paragraph(_pdf_escape(scenario_obj.description[:1200]), body))

    story.append(Paragraph(_pdf_escape("二、专项核查清单（摘要）"), h2))
    for section in payload.get("sections", []):
        story.append(Paragraph(_pdf_escape(section.get("dimension_name", "")), h2))
        for item in section.get("items", [])[:6]:
            line = f"[{item.get('code')}] {item.get('title')}（{item.get('priority')}）"
            story.append(Paragraph(_pdf_escape(line), body))

    story.append(Paragraph(_pdf_escape("三、法源绑定摘要"), h2))
    for section in sections_legal:
        for item in section.get("items", []):
            hits = item.get("legal_hits") or []
            if not hits:
                continue
            hit = hits[0]
            line = (
                f"{item.get('code')} — {hit.get('title_zh') or hit.get('title_pt')} "
                f"(匹配度 {hit.get('match_score', 0)})"
            )
            story.append(Paragraph(_pdf_escape(line), body))

    if brief:
        story.append(Paragraph(_pdf_escape("四、双语简报摘要"), h2))
        story.append(Paragraph(_pdf_escape(brief.get("summary_zh", "")[:800]), body))
        story.append(Paragraph(_pdf_escape(brief.get("summary_pt", "")[:800]), body))

    if review and review.get("items"):
        story.append(Paragraph(_pdf_escape("五、法务复核记录"), h2))
        story.append(
            Paragraph(
                _pdf_escape(f"复核人：{review.get('reviewer_name', '')} · 状态：{review.get('status', '')}"),
                body,
            )
        )

    story.append(Paragraph(_pdf_escape("六、免责声明"), h2))
    for para in DISCLAIMER_FULL_TEXT.split("\n\n")[:4]:
        story.append(Paragraph(_pdf_escape(para), body))

    doc.build(story)
    filename = safe_export_filename(scenario_obj.project_name, "协查底稿.pdf")
    return buf.getvalue(), filename
