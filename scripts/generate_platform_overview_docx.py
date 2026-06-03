#!/usr/bin/env python3
"""Generate Word overview: Vela platform business & legal capabilities + implementation."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Vela平台_业务与法务能力及实现说明.docx"


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def build() -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Vela 出海法务平台\n业务端与法务端能力及实现说明")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"生成日期：{date.today().isoformat()}  ·  法域：巴西  ·  演示场景：BYD 坎皮纳斯")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    _heading(doc, "一、平台概述")
    _para(
        doc,
        "Vela 出海法务平台面向企业出海投资合规协查，主链路为：中文场景输入 → 专项核查清单 → "
        "多源法条检索（RAG）→ 70 分门控双语简报 → 法务复核 → 定稿导出 Word/PDF。",
    )
    _para(doc, "技术栈：前端 Vue 3 + Pinia；后端 FastAPI + SQLAlchemy；开发库 SQLite；法源索引 Chroma；可选 LLM 润色。", bold=False)

    _heading(doc, "二、角色与权限")
    _bullets(
        doc,
        [
            "业务账号（business）：提交协查、跟踪进度、补充材料、业务侧回收站。",
            "法务账号（legal/admin）：复核定稿、导出底稿、法源维护、法务侧回收站。",
            "后端通过 require_role / get_current_legal_user 控制 API；前端通过 Vue Router meta（legalOnly / businessOnly）限制页面。",
            "演示账户：biz@demo.vela / legal@demo.vela，密码 Demo1234!。",
        ],
    )

    _heading(doc, "三、业务端能力与实现")
    _heading(doc, "3.1 工作台", 2)
    _bullets(
        doc,
        [
            "提交新协查（折叠面板）：ScenarioCreateView.vue → POST /scenarios/generate-and-submit。",
            "我的协查进度（折叠）：DashboardView.vue 按 待补充 / 进行中 / 已完成 三组折叠展示。",
            "回收站（折叠）：POST /archive 写 business_archived_at；POST /unarchive 恢复。",
        ],
    )

    _heading(doc, "3.2 提交协查", 2)
    _bullets(
        doc,
        [
            "手动填写项目信息，或上传 .txt / .md / .docx / .pdf（≤2MB）预填：POST /scenarios/extract-document。",
            "一键流程（scenario_pipeline.run_generate_and_submit）：生成清单 → RAG 法条 → 双语简报 → status=pending_legal_review。",
            "清单：rule_engine.generate_checklist()；法条：legal_rag.retrieve_for_checklist()；简报：brief_generator.generate_brief()。",
            "BYD 演示：POST /scenarios/demo/generate-and-submit。",
        ],
    )

    _heading(doc, "3.3 跟踪与协作", 2)
    _bullets(
        doc,
        [
            "查看进度：BusinessProgressView.vue，读取 scenario + business_feedback_from_review()。",
            "查看清单/简报/AI 抽取：ChecklistView / BriefView / BusinessExtractView（只读）。",
            "法务退回后补充：ScenarioCreateView edit 模式 → POST /revise-and-resubmit → 同一项目重提。",
            "分组规则：needs_revision → 待补充；progress_status=finalized → 已完成；其余 → 进行中。",
        ],
    )

    _heading(doc, "四、法务端能力与实现")
    _heading(doc, "4.1 工作台", 2)
    _bullets(
        doc,
        [
            "全部协查场景（折叠，置顶）：四组 待复核 / 复核中 / 待业务补充 / 已完成；前两组按 review_priority 排序。",
            "回收站（折叠，第二位）：DELETE /scenarios/{id} 软删 legal_deleted_at；POST /restore 恢复。",
            "法源语料维护：LegalCorpusView.vue + legal.py corpus CRUD + reindex。",
            "法务演示：POST /scenarios/demo/sample 一键生成完整样本直达复核。",
            "合规审查维度 / 系统状态：规则包只读展示与健康检查。",
        ],
    )

    _heading(doc, "4.2 复核与定稿", 2)
    _bullets(
        doc,
        [
            "ReviewView.vue：按合规维度分组；展开行内嵌简报片段；逐条确认/驳回/批注。",
            "POST /review/init 初始化；PATCH /review/items/{code} 更新条目；POST /review/approve-all 批量确认。",
            "POST /review/return-to-business：至少一条驳回 → status=returned_for_revision。",
            "POST /review/finalize：全部条目处理完 → review_approved/rejected/partial。",
            "复核数据存于 checklist.payload.review（JSON）。",
        ],
    )

    _heading(doc, "4.3 导出与法源", 2)
    _bullets(
        doc,
        [
            "定稿后 GET /export/docx（法学院《法律研究意见书》）或 /export/pdf（legacy 协查底稿）。",
            "export_service.build_law_school_docx() / build_sample_pdf()。",
            "法源语料维护后须 POST /corpus/reindex，新协查才引用最新索引。",
            "法规监测：GET/POST /legal/monitor/scan（MVP，本地语料变更提醒）。",
        ],
    )

    _heading(doc, "五、数据模型与状态机")
    _para(doc, "核心表：investigation_scenarios + compliance_checklists.payload（JSON）。")
    _bullets(
        doc,
        [
            "payload 含：sections、sections_with_legal、brief、review、document_extract。",
            "scenario.status 驱动流程；list_scenarios 时 scenario_summary_meta() 映射 progress_status 供前端分组。",
            "pending_legal_review → submitted；review_in_progress → in_review；returned_for_revision → needs_revision；review_* → finalized。",
            "业务回收站：business_archived_at；法务回收站：legal_deleted_at（互不影响业务访问）。",
        ],
    )

    _heading(doc, "六、协作流程")
    _para(
        doc,
        "业务 generate-and-submit → 法务 review/init → 逐条复核 → 定稿 export；"
        "若退回：return-to-business → 业务 revise-and-resubmit → 再次 pending_legal_review。",
    )

    _heading(doc, "七、本地启动")
    _bullets(
        doc,
        [
            "项目根目录执行：./scripts/start.sh",
            "前端：http://127.0.0.1:5173（须带端口 :5173）",
            "后端 API：http://127.0.0.1:8000/docs",
            "可选 LLM：backend/.env 配置 QWEN_API_KEY 或 DEEPSEEK_API_KEY，LLM_POLISH_ENABLED=true",
        ],
    )

    _heading(doc, "八、主要代码路径速查")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "模块"
    hdr[1].text = "后端"
    hdr[2].text = "前端"
    rows = [
        ("提交协查", "scenario_pipeline.py / scenarios.py", "ScenarioCreateView.vue"),
        ("工作台", "scenarios.py list + archive", "DashboardView.vue"),
        ("法务复核", "review_service.py / scenarios.py", "ReviewView.vue"),
        ("清单/简报", "rule_engine / brief_generator", "ChecklistView / BriefView"),
        ("法源维护", "legal.py / legal_corpus_service", "LegalCorpusView.vue"),
        ("导出", "export_service.py", "ReviewView 导出按钮"),
        ("权限", "core/roles.py / deps.py", "router/index.ts + stores/auth.ts"),
    ]
    for mod, be, fe in rows:
        row = table.add_row().cells
        row[0].text = mod
        row[1].text = be
        row[2].text = fe

    doc.add_paragraph()
    _para(
        doc,
        "免责声明：本文档为 Vela 平台能力与实现说明，供内部培训与演示使用；"
        "不构成正式法律意见。定稿协查结论以法务复核后导出的 Word/PDF 为准。",
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
