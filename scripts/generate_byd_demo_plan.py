#!/usr/bin/env python3
"""Generate BYD Campinas demo investment plan PDF + DOCX for AI extract upload."""

from __future__ import annotations

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

TITLE = "BYD 坎皮纳斯新能源工厂投资方案（演示）"

SECTIONS: list[tuple[str, str]] = [
    (
        "一、项目概况",
        "项目名称：BYD 坎皮纳斯新能源工厂（演示案例）\n"
        "项目地点：巴西 · 圣保罗州 · 坎皮纳斯市（Campinas）\n"
        "行业方向：新能源制造 · 绿地建厂（Greenfield）\n"
        "投资主体：中资母公司通过巴西全资子公司投资，100% 外资",
    ),
    (
        "二、投资结构与用工",
        "投资结构：中资母公司通过巴西全资子公司投资，100% 外资。\n"
        "雇员规模：项目将为当地创造约 450 名新的就业岗位，含生产、工程及本地化管理岗位。\n"
        "外派安排：初期拟派驻少量中方技术与管理骨干，其余岗位优先本地招聘。",
    ),
    (
        "三、生产设施规划",
        "产能规划：首年具备最高 1000 辆电动大巴及其配套磷酸铁锂电池包的年生产能力；"
        "并规划太阳能电池板与储能系统产线。\n"
        "厂房设施：制造设施占地约 32,000 平方米，另配套约 20,000 平方米厂房；"
        "同时计划设立光伏、智能电网与 LED 照明业务的研发中心。",
    ),
    (
        "四、项目描述",
        "巴西圣保罗州的坎皮纳斯市（Campinas）已被选为比亚迪最新工厂的所在地，"
        "该项目将为当地创造约 450 个新的就业岗位。该工厂预计于 2020 年投产，"
        "不仅将启动南美洲唯一一款长续航纯电动公交车的制造与组装业务，"
        "还将生产首款具备极高防火安全性且可回收的磷酸铁锂电池包。"
        "在投产的第一年，该工厂将具备最高 1000 辆电动大巴及其所有配套电池的年生产能力。"
        "除了大巴和电池，还计划生产太阳能电池板和储能系统。"
        "除了占地 32,000 平方米和 20,000 平方米的制造设施外，"
        "还计划为光伏、智能电网和 LED 照明业务开设研发中心。",
    ),
    (
        "五、时间线",
        "董事会/投委会日期：2013 年 12 月 3 日\n"
        "预计开工：2015 年 6 月\n"
        "预计投产：2020 年 10 月",
    ),
    (
        "六、合规关注（业务说明）",
        "本项目涉及本地用工（CLT）、100% 外资设厂、州级 ICMS/PIS/COFINS 税制激励评估、"
        "环评与工业许可、电池与客车行业准入等合规事项，需开展专项协查。",
    ),
    (
        "七、备注",
        "本文件为 Vela 平台 BYD 坎皮纳斯演示预设场景配套材料，仅供 AI 抽取与协查流程体验，"
        "不代表真实投资决策文件。",
    ),
]


def _pdf_style(name: str, **kwargs) -> ParagraphStyle:
    base = {"fontName": "STSong-Light", "fontSize": 11, "leading": 18}
    base.update(kwargs)
    return ParagraphStyle(name, **base)


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(path: Path) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    story = []
    story.append(Paragraph(_escape(TITLE), _pdf_style("Title", fontSize=16, leading=22, spaceAfter=12)))
    story.append(
        Paragraph(
            _escape("文档类型：对外投资设厂方案摘要 · 密级：内部演示"),
            _pdf_style("Sub", fontSize=10, textColor="#555555", spaceAfter=16),
        )
    )
    for heading, body in SECTIONS:
        story.append(Paragraph(_escape(heading), _pdf_style("H2", fontSize=13, leading=20, spaceBefore=10, spaceAfter=6)))
        for para in body.split("\n"):
            story.append(Paragraph(_escape(para), _pdf_style("Body")))
        story.append(Spacer(1, 8))
    doc.build(story)


def build_docx(path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal.font.size = Pt(11)

    title = doc.add_heading(TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("文档类型：对外投资设厂方案摘要 · 密级：内部演示").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for heading, body in SECTIONS:
        doc.add_heading(heading, level=1)
        for para in body.split("\n"):
            doc.add_paragraph(para)

    doc.save(path)


def main() -> None:
    pdf_path = ROOT / "BYD坎皮纳斯_投资方案_演示.pdf"
    docx_path = ROOT / "BYD坎皮纳斯_投资方案_演示.docx"
    build_pdf(pdf_path)
    build_docx(docx_path)
    print(f"Wrote {pdf_path}")
    print(f"Wrote {docx_path}")
    print("")
    print("AI 上传体验：登录 biz@demo.vela → 提交新协查 → 上传 .pdf 或 .docx 文件")
    print("（平台抽取支持 .txt / .md / .docx / .pdf）")


if __name__ == "__main__":
    main()
